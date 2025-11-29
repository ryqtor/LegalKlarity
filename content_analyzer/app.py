import io
import re
import pdfplumber
import docx
import pytesseract
import fitz
from PIL import Image
from flask import Flask, request, jsonify, send_file
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
import json
import os
from datetime import datetime
import textwrap

# Try to import Google AI Studio SDK
try:
    import google.generativeai as genai
    from dotenv import load_dotenv
    load_dotenv()
    
    API_KEY = os.environ.get("GEMINI_API_KEY", "").strip()
    
    if API_KEY:
        genai.configure(api_key=API_KEY)
        # Use gemini-flash-latest as gemini-1.5-flash is not available
        model = genai.GenerativeModel("gemini-flash-latest")
        AI_MODE = "STUDIO"
        print("="*60)
        print("AI MODE ENABLED - Using Google AI Studio")
        print("="*60)
    else:
        AI_MODE = "NONE"
        model = None
        print("No GEMINI_API_KEY found - using fallback analysis")

except ImportError:
    AI_MODE = "NONE"
    model = None
    print("Google AI SDK not available - using fallback analysis")
except Exception as e:
    AI_MODE = "NONE"
    model = None
    print(f"AI initialization failed: {e}")

# Flask app
app = Flask(__name__)

# Google Cloud configuration
GOOGLE_CLOUD_PROJECT = os.environ.get("GOOGLE_CLOUD_PROJECT", "your-google-cloud-project-id")
GOOGLE_CLOUD_LOCATION = os.environ.get("GOOGLE_CLOUD_LOCATION", "us-central1")

# Section cues to check for agreements
POSITIVE_LABELS = [
    "agreement", "legal contract", "rental agreement", "lease agreement",
    "service agreement", "tenant-landlord agreement", "terms and conditions",
    "offer letter", "internship agreement", "employment contract",
    "student agreement", "job offer", "internship terms"
]

SECTION_CUES = [
    "agreement", "security deposit", "rental period", "payment terms",
    "termination", "arbitration", "jurisdiction",
    "witness", "signatory", "governing law", "parties", "definitions",
    "probation period", "internship duration", "performance", 
    "salary", "compensation", "notice period", "work expectations",
    "attendance", "leaves", "certificate", "offer letter"
]

# Helpers
def safe_join_text(parts):
    return "\n".join([p for p in parts if p])

def chunk_text(text, max_words=300, max_chunks=10):
    """
    Split text into chunks of specified word count
    """
    words = text.split()
    chunks = [" ".join(words[i:i + max_words]) for i in range(0, len(words), max_words)]
    return chunks[:max_chunks]

def heuristic_score(text):
    t = (text or "").lower()
    found = sum(1 for k in SECTION_CUES if re.search(r"\b" + re.escape(k) + r"\b", t))
    return found / max(1, len(SECTION_CUES))

def classify_agreement(text):
    details = {
        "chunks": 0,
        "votes": 0,
        "vote_ratio": 0.0,
        "heuristic": 0.0,
        "avg_chunk_score": 0.0,
        "reason": ""
    }
    if not text.strip():
        details["reason"] = "empty_text"
        return False, details
    chunks = chunk_text(text, max_words=300, max_chunks=10)
    details["chunks"] = len(chunks)
    votes, per_chunk_scores = 0, []
    CHUNK_THRESHOLD = 0.5
    for ch in chunks:
        # Simple keyword-based classification instead of ML model
        score = heuristic_score(ch)
        per_chunk_scores.append(score)
        if score >= CHUNK_THRESHOLD:
            votes += 1
    ratio = votes / len(chunks)
    heur = heuristic_score(text)
    details.update({
        "votes": votes,
        "vote_ratio": round(ratio, 3),
        "heuristic": round(heur, 3),
        "avg_chunk_score": round(sum(per_chunk_scores) / max(1, len(per_chunk_scores)), 3)
    })
    accept = (ratio >= 0.4) or (heur >= 0.4)
    if not accept:
        details["reason"] = "low_confidence"
    return accept, details

# Document type detection
def detect_document_type(text):
    """
    Enhanced document type detection
    """
    text_lower = text.lower()
    
    # Document type patterns
    patterns = {
        "rental agreement": ["rent", "lease", "tenant", "landlord", "security deposit"],
        "employment contract": ["employment", "employee", "employer", "salary", "position"],
        "service agreement": ["service", "provider", "client", "deliverable"],
        "loan agreement": ["loan", "borrower", "lender", "interest rate"],
        "nda": ["confidential", "non-disclosure", "secrecy"],
        "purchase agreement": ["purchase", "buy", "sell", "buyer", "seller"],
        "internship agreement": ["internship", "intern", "supervisor", "internship period"]
    }
    
    # Score each document type
    scores = {}
    for doc_type, keywords in patterns.items():
        score = sum(1 for keyword in keywords if keyword in text_lower)
        scores[doc_type] = score
    
    # Return highest scoring document type
    if scores:
        best_match = max(scores.items(), key=lambda x: x[1])
        if best_match[1] > 0:
            return best_match[0]
    
    return "general legal document"

# Fallback analysis function
def create_fallback_analysis(text, document_type):
    """
    Create basic analysis when AI analysis fails
    """
    # Extract first few sentences for summary
    sentences = text.split('.')
    summary = '. '.join(sentences[:3]) + '.' if len(sentences) > 3 else text[:500]
    
    return {
        "summary": summary,
        "key_terms": [],
        "main_clauses": [],
        "critical_dates": [],
        "parties": [],
        "jurisdiction": "Not analyzed",
        "obligations": [],
        "risks": [],
        "recommendations": ["Have a legal professional review this document"],
        "missing_clauses": [],
        "compliance_issues": [],
        "next_steps": ["Review document with legal counsel"]
    }

# Enhanced document analysis function
def analyze_legal_document(text, document_type=None):
    """
    Comprehensive legal document analysis using Gemini AI
    """
    
    # Auto-detect document type if not provided
    if not document_type:
        document_type = detect_document_type(text)
    
    # If AI is not available, use fallback analysis
    if AI_MODE == "NONE" or model is None:
        print("Using fallback analysis - AI not available")
        return create_fallback_analysis(text, document_type)
    
    # Enhanced prompt engineering for comprehensive analysis
    prompt = f"""
    Analyze the following {document_type or 'legal document'} and provide a comprehensive analysis.
    Return ONLY valid JSON that strictly matches this schema:
    
    {{
        "summary": "Brief 2-3 sentence overview of the entire document",
        "key_terms": [
            {{
                "term": "Defined term",
                "definition": "Clear definition from the document"
            }}
        ],
        "main_clauses": [
            {{
                "name": "Clause name/title",
                "description": "Brief description of what this clause covers"
            }}
        ],
        "critical_dates": [
            {{
                "date": "YYYY-MM-DD or date range",
                "event": "What happens on this date"
            }}
        ],
        "parties": [
            {{
                "name": "Party name",
                "role": "Their role in the agreement"
            }}
        ],
        "jurisdiction": "Governing law and jurisdiction information",
        "obligations": [
            {{
                "party": "Which party",
                "responsibility": "What they must do"
            }}
        ],
        "risks": [
            {{
                "risk": "Identified risk",
                "severity": "high/medium/low",
                "description": "Explanation of the risk"
            }}
        ],
        "recommendations": [
            "Actionable recommendation to address identified issues"
        ],
        "missing_clauses": [
            {{
                "clause": "Missing clause name",
                "importance": "Why it's important"
            }}
        ],
        "compliance_issues": [
            {{
                "issue": "Compliance concern",
                "regulation": "Relevant law/regulation (if identifiable)"
            }}
        ],
        "next_steps": [
            "Action item that should be taken next"
        ]
    }}
    
    Document Text:
    {text[:50000]}  # Limit to prevent token overflow
    
    Rules:
    - If information is not found, return an empty string ("") or empty list ([]).
    - Do not include explanations outside the JSON.
    - Keep responses concise and accessible for non-lawyers.
    - While generating each field, add more and more explanation and context to each field to ensure deep analysis.
    - Strictly return JSON with the specified fields and no additional fields.
    """
    
    try:
        if AI_MODE == "NONE" or model is None:
            raise Exception("No AI model initialized")

        # Generate response
        response = model.generate_content(
            prompt,
            generation_config={
                "temperature": 0.4,
                "top_p": 0.8,
                "top_k": 40,
                "max_output_tokens": 8192,
            }
        )
        
        # Clean response text (remove markdown code blocks)
        response_text = response.text.strip()
        if response_text.startswith("```json"):
            response_text = response_text[7:]
        elif response_text.startswith("```"):
            response_text = response_text[3:]
        
        if response_text.endswith("```"):
            response_text = response_text[:-3]
            
        response_text = response_text.strip()
        
        # Parse and validate JSON response
        analysis = json.loads(response_text)
        return analysis
        
    except json.JSONDecodeError as e:
        print(f"JSON parsing failed: {e}")
        # Fallback to basic analysis if JSON parsing fails
        return create_fallback_analysis(text, document_type)
    except Exception as e:
        print(f"Analysis failed: {e}")
        # Return error structure
        return {
            "error": f"Analysis failed: {str(e)}",
            "summary": "Document analysis could not be completed due to technical issues.",
            "key_terms": [],
            "main_clauses": [],
            "critical_dates": [],
            "parties": [],
            "jurisdiction": "Not available",
            "obligations": [],
            "risks": [],
            "recommendations": [],
            "missing_clauses": [],
            "compliance_issues": [],
            "next_steps": []
        }

# Enhanced Flask route for document analysis
@app.route("/enhanced_analysis", methods=["POST"])
def enhanced_document_analysis():
    """
    Enhanced document analysis endpoint
    """
    print("Received request to enhanced_analysis endpoint")
    
    if "file" not in request.files:
        print("No file uploaded")
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files["file"]
    print(f"Received file: {file.filename}")
    
    if file.filename == "":
        print("No file selected")
        return jsonify({"error": "No file selected"}), 400

    # Extract text (using existing functions)
    filename = file.filename.lower()
    text = ""
    
    try:
        if filename.endswith(".pdf"):
            print("Processing PDF file")
            text = extract_pdf(file.stream)
        elif filename.endswith(".docx"):
            print("Processing DOCX file")
            text = extract_docx(file.stream)
        elif filename.endswith((".png", ".jpg", ".jpeg")):
            print("Processing image file")
            text = extract_image(file.stream)
        else:
            print(f"Unsupported file type: {filename}")
            return jsonify({"error": "Unsupported file type"}), 400
        
        print(f"Extracted text length: {len(text)}")
        
        # Check if it's a valid agreement (using existing function)
        is_ok, details = classify_agreement(text)
        print(f"Classification result: {is_ok}, Details: {details}")
        
        if not is_ok:
            print("Warning: Low confidence classification, proceeding anyway")
        
        # Perform enhanced analysis
        print("Performing enhanced analysis")
        analysis = analyze_legal_document(text)
        print(f"Analysis completed: {analysis.get('summary', 'No summary')[:100]}...")
        
        return jsonify({
            "filename": file.filename,
            "extracted_text": text,
            "analysis": analysis,
            "timestamp": datetime.now().isoformat()
        })
    except Exception as e:
        print(f"Error in enhanced_document_analysis: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"Internal server error: {str(e)}"}), 500

# File extraction functions
def extract_pdf(file_stream):
    try:
        file_stream.seek(0)
        with pdfplumber.open(file_stream) as pdf:
            return safe_join_text([p.extract_text() for p in pdf.pages])
    except Exception as e:
        print(f"PDF extract error: {e}")
        try:
            file_stream.seek(0)
            doc = fitz.open(stream=file_stream.read(), filetype="pdf")
            texts = []
            for p in doc:
                pix = p.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                texts.append(pytesseract.image_to_string(img))
            return "\n".join(texts)
        except Exception as e2:
            print(f"PDF OCR error: {e2}")
            return ""

def extract_docx(file_stream):
    try:
        file_stream.seek(0)
        doc = docx.Document(io.BytesIO(file_stream.read()))
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    except Exception as e:
        print(f"DOCX extract error: {e}")
        return ""

def extract_image(file_stream):
    try:
        file_stream.seek(0)
        img = Image.open(file_stream).convert("RGB")
        return pytesseract.image_to_string(img)
    except Exception as e:
        print(f"Image extract error: {e}")
        return ""

# Routes
@app.route("/active", methods=["GET"])
def active():
    return "active"

@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    text = request.form.get("text", "")
    output = io.BytesIO()
    doc = SimpleDocTemplate(output)
    styles = getSampleStyleSheet()
    story = [Paragraph(line, styles["Normal"]) for line in text.split("\n")]
    doc.build(story)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.pdf")

@app.route("/export/docx", methods=["POST"])
def export_docx():
    text = request.form.get("text", "")
    output = io.BytesIO()
    d = docx.Document()
    for line in text.split("\n"):
        d.add_paragraph(line)
    d.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.docx")

if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port)